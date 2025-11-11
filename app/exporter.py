"""
DOCX файл үүсгэх модуль
"""

from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
from pathlib import Path


def export_enhanced_protocol(protocol: dict) -> str:
    """
    Протоколыг DOCX файл болгох
    
    Args:
        protocol: Протоколын мэдээлэл
            - title: Гарчиг
            - date: Огноо
            - participants: Оролцогчид (list)
            - body: Үндсэн агуулга
            - action_items: Ажил үүргүүд (list of dict)
    
    Returns:
        Үүсгэсэн файлын нэр
    
    Raises:
        RuntimeError: Файл үүсгэхэд алдаа гарвал
    """
    try:
        doc = Document()
        
        # 1. Толгой
        _add_heading(doc, protocol)
        
        # 2. Мета мэдээлэл
        _add_metadata(doc, protocol)
        
        # 3. Үндсэн агуулга
        _add_body(doc, protocol)
        
        # 4. Ажил үүргийн хүснэгт
        if protocol.get('action_items'):
            _add_action_items_table(doc, protocol['action_items'])
        
        # 5. Footer
        _add_footer(doc)
        
        # 6. Хадгалах
        filename = f"protocol_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        
        # Файл үүссэн эсэхийг шалгах
        if not Path(filename).exists():
            raise RuntimeError(f"Файл үүсгэж чадсангүй: {filename}")
        
        return filename
        
    except Exception as e:
        raise RuntimeError(f"DOCX export алдаа: {str(e)}")


def _add_heading(doc: Document, protocol: dict):
    """Толгой хэсэг нэмэх"""
    title = doc.add_heading(protocol.get('title', 'Протокол'), level=1)
    
    # Толгой форматлах
    for run in title.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 120)


def _add_metadata(doc: Document, protocol: dict):
    """Мета мэдээлэл нэмэх"""
    doc.add_paragraph()
    
    # Огноо
    date_para = doc.add_paragraph()
    date_run = date_para.add_run("Огноо: ")
    date_run.bold = True
    date_para.add_run(protocol.get('date', ''))
    
    # Оролцогчид
    participants_para = doc.add_paragraph()
    participants_run = participants_para.add_run("Оролцогчид: ")
    participants_run.bold = True
    
    participants = protocol.get('participants', [])
    if isinstance(participants, list):
        participants_text = ', '.join(participants)
    else:
        participants_text = str(participants)
    
    participants_para.add_run(participants_text)
    
    # Тусгаарлагч
    doc.add_paragraph("_" * 70)


def _add_body(doc: Document, protocol: dict):
    """Үндсэн агуулга нэмэх"""
    doc.add_heading("Хэлэлцсэн асуудал", level=2)
    
    body_text = protocol.get('body', '')
    
    # Paragraph болгон хуваах
    paragraphs = body_text.split('\n')
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            para = doc.add_paragraph(para_text)
            
            # Формат тохируулах
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)


def _add_action_items_table(doc: Document, action_items: list):
    """Ажил үүргийн хүснэгт нэмэх"""
    doc.add_heading("Ажил үүрэг ба шийдвэрүүд", level=2)
    
    # Хүснэгт үүсгэх
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    headers = ['Хариуцагч', 'Ажил үүрэг', 'Хугацаа', 'Төрөл']
    
    for idx, header_text in enumerate(headers):
        cell = hdr_cells[idx]
        cell.text = header_text
        
        # Header форматлах
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Data rows
    for action in action_items:
        row = table.add_row().cells
        
        row[0].text = action.get('who', '')
        row[1].text = action.get('action', '')
        row[2].text = action.get('due', 'Тодорхойгүй')
        
        # Төрөл (Монголоор)
        action_type = action.get('type', 'action')
        row[3].text = 'Шийдвэр' if action_type == 'decision' else 'Ажил үүрэг'
        
        # Row форматлах
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    # Summary
    doc.add_paragraph()
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(f"Нийт: {len(action_items)} ажил үүрэг/шийдвэр")
    summary_run.italic = True
    summary_run.font.size = Pt(9)


def _add_footer(doc: Document):
    """Footer нэмэх"""
    doc.add_paragraph()
    doc.add_paragraph("_" * 70)
    
    footer_para = doc.add_paragraph()
    footer_text = f"Протокол автоматаар үүсгэгдсэн | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    footer_run = footer_para.add_run(footer_text)
    
    # Footer форматлах
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True
    
    # Төвд байрлуулах
    footer_para.alignment = 1  # CENTER